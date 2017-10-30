#!/usr/bin/env python
# -*- coding: utf-8 -*-

import os
import os.path
import argparse
from docx import Document
from docx.shared import Inches

__author__ = 'heshed'

def convert_to_docx(input_dir, output_dir):
    print '--- [MS-WORD 제작 시작]'

    document_size = 1000
    document_pool = {}

    if not os.path.isdir(output_dir):
        print '    create dir ', output_dir
        os.mkdir(output_dir)

    for item in os.listdir(input_dir):
        item_encoding = encoding(item)
        if 'euc-kr' in item_encoding:
            item = unicode(item, 'euc-kr').encode('utf-8')

        sub_dir = os.path.join(input_dir, item)

        if os.path.isdir(sub_dir):
            index = 1
            file_count = 0

            def numeric_diff(a, b):
                ai = int(os.path.splitext(os.path.basename(a))[0])
                bi = int(os.path.splitext(os.path.basename(b))[0])
                return ai - bi

            news_sorted = sorted(os.listdir(sub_dir), cmp=numeric_diff)
            for news in news_sorted:
                file_count += 1
                if file_count % document_size == 0:
                    index += 1

                file_name = os.path.join(output_dir, '{}-{}.docx'.format(item, index))

                if file_count % 100 == 0:
                    print '.',
                if file_count % document_size == 0:
                    print
                    print file_name

                document = get_document(file_name, document_pool)

                file_path = os.path.join(sub_dir, news)
                if os.path.isfile(file_path):
                    title = '언론사: {}-{}-{}'.format(item, index, news)
                    title_encoding = encoding(title)
                    document.add_heading(title.decode(title_encoding), 0)
                    content = read_file(file_path)
                    cont_encoding = encoding(content)
                    document.add_paragraph(content.decode(cont_encoding))

    print '--- [MS-WORD 쓰기 시작]'

    for key, val in document_pool.items():
        print key
        val.save(key)

    print '--- [MS-WORD 쓰기 끝]'


def get_document(path, pool):
    if not path in pool:
        pool[path] = Document()
    return pool.get(path)


def read_file(path):
    rf = file(path)
    data = rf.read()
    rf.close()
    return data


def encoding(s):
    try:
        s.decode('utf-8')
        return 'utf-8'
    except UnicodeError:
        pass

    try:
        s.decode('euc-kr')
        return 'euc-kr'
    except UnicodeError:
        pass

    return ''

if __name__ == '__main__':
    default_input_path = '결과-상세'
    default_output_path = 'msword'

    parser = argparse.ArgumentParser(description="상세 뉴스 데이터를 MS Word로 생성합니다.")
    parser.add_argument("-i", default=default_input_path, help="input directory: ex) " + default_input_path)
    parser.add_argument("-o", default=default_output_path, help="output directory: ex) " + default_output_path)
    args = parser.parse_args()

    convert_to_docx(args.i, args.o)
